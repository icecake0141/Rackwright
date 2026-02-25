"""
Copyright 2026 Rackwright Contributors
SPDX-License-Identifier: Apache-2.0

This file was created or modified with the assistance of an AI (Large Language Model).
Review required for correctness, security, and licensing.
"""

from __future__ import annotations

import hashlib
import json
import os
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from jinja2 import Template
from openpyxl import Workbook
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from rackwright.change_logging import change_log_context
from rackwright.models import (
    ArtifactFile,
    ArtifactVersion,
    Cabling,
    Device,
    PowerCabling,
    Project,
    Rack,
    Room,
    Row,
    SectionApplicationRule,
    SectionSnapshot,
    Site,
    TemplateSetSnapshot,
)
from rackwright.view_builders import cablings_for_project

VALID_MODES = {"all", "word", "excel", "images"}


@dataclass
class GenerationResult:
    artifact_version: ArtifactVersion
    skipped: bool


def _data_dir() -> Path:
    base = os.environ.get("RACKWRIGHT_DATA_DIR", "./data")
    return Path(base)


def _load_project(session: Session, project_id: int) -> Project:
    project = session.get(Project, project_id)
    if project is None:
        raise ValueError("Project not found")
    return project


def _serialize_rows(rows: list[Any], fields: list[str]) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    for row in rows:
        result.append({field: getattr(row, field) for field in fields})
    return result


def _compute_fingerprint(
    session: Session, project_id: int, mode: str, base_version_id: int | None
) -> str:
    project = _load_project(session, project_id)
    racks = (
        session.execute(
            select(Rack).where(Rack.project_id == project_id).order_by(Rack.id)
        )
        .scalars()
        .all()
    )
    devices = (
        session.execute(
            select(Device).where(Device.project_id == project_id).order_by(Device.id)
        )
        .scalars()
        .all()
    )
    cablings = (
        session.execute(
            select(Cabling).where(Cabling.project_id == project_id).order_by(Cabling.id)
        )
        .scalars()
        .all()
    )
    power = (
        session.execute(
            select(PowerCabling)
            .where(PowerCabling.project_id == project_id)
            .order_by(PowerCabling.id)
        )
        .scalars()
        .all()
    )

    snapshots = (
        session.execute(
            select(SectionSnapshot)
            .join(
                TemplateSetSnapshot,
                TemplateSetSnapshot.id == SectionSnapshot.template_set_snapshot_id,
            )
            .where(TemplateSetSnapshot.project_id == project_id)
            .order_by(
                SectionSnapshot.category,
                SectionSnapshot.section_order,
                SectionSnapshot.id,
            )
        )
        .scalars()
        .all()
    )
    rules = (
        session.execute(
            select(SectionApplicationRule)
            .where(SectionApplicationRule.project_id == project_id)
            .order_by(SectionApplicationRule.section_snapshot_id)
        )
        .scalars()
        .all()
    )

    payload = {
        "project": {
            "id": project.id,
            "name": project.name,
            "owner": project.owner,
            "notes": project.notes,
            "revision": project.revision,
        },
        "racks": _serialize_rows(racks, ["id", "name", "rack_height_u", "row_id"]),
        "devices": _serialize_rows(
            devices,
            [
                "id",
                "name",
                "role",
                "model",
                "serial",
                "power_watts",
                "rack_id",
                "ru_start",
                "ru_size",
                "orientation",
            ],
        ),
        "cablings": _serialize_rows(
            cablings,
            [
                "id",
                "normalized_key",
                "a_device",
                "a_port",
                "a_port_type",
                "b_device",
                "b_port",
                "b_port_type",
                "cable_type",
                "label",
            ],
        ),
        "power_cablings": _serialize_rows(
            power,
            [
                "id",
                "normalized_key",
                "a_device",
                "a_port",
                "a_port_type",
                "b_device",
                "b_port",
                "b_port_type",
                "bank",
                "outlet",
                "label",
            ],
        ),
        "sections": _serialize_rows(
            snapshots,
            [
                "id",
                "target_type",
                "category",
                "section_order",
                "output_targets",
                "applicable_roles",
                "text",
            ],
        ),
        "rules": _serialize_rows(
            rules, ["section_snapshot_id", "enabled", "filters_json"]
        ),
        "mode": mode,
        "base_version_id": base_version_id,
    }

    encoded = json.dumps(payload, sort_keys=True, ensure_ascii=False).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _next_version_number(session: Session, project_id: int) -> int:
    current = session.execute(
        select(func.max(ArtifactVersion.version_number)).where(
            ArtifactVersion.project_id == project_id
        )
    ).scalar()
    return (current or 0) + 1


def _section_rules_map(
    session: Session, project_id: int
) -> dict[int, SectionApplicationRule]:
    rules = (
        session.execute(
            select(SectionApplicationRule).where(
                SectionApplicationRule.project_id == project_id
            )
        )
        .scalars()
        .all()
    )
    return {rule.section_snapshot_id: rule for rule in rules}


def _parse_filter_values(filters: dict[str, Any], key: str) -> set[str]:
    value = filters.get(key)
    if value is None:
        return set()
    if isinstance(value, list):
        return {str(item).strip().lower() for item in value if str(item).strip()}
    return {str(value).strip().lower()} if str(value).strip() else set()


def _rule_filters(rule: SectionApplicationRule | None) -> dict[str, Any]:
    if rule is None or not rule.filters_json:
        return {}
    payload = json.loads(rule.filters_json)
    if isinstance(payload, dict):
        return payload
    return {}


def _target_objects_for_section(
    session: Session,
    project: Project,
    section: SectionSnapshot,
    rule: SectionApplicationRule | None,
) -> list[Any]:
    if rule is not None and not rule.enabled:
        return []

    filters = _rule_filters(rule)
    role_values = _parse_filter_values(filters, "role")
    rack_scope_values = _parse_filter_values(filters, "rack_scope")

    if section.target_type == "Project":
        return [project]

    racks = (
        session.execute(
            select(Rack).where(Rack.project_id == project.id).order_by(Rack.id)
        )
        .scalars()
        .all()
    )
    rack_id_to_name = {rack.id: rack.name for rack in racks}

    if section.target_type == "Rack":
        if not rack_scope_values:
            return list(racks)
        return [
            rack
            for rack in racks
            if str(rack.id).lower() in rack_scope_values
            or rack.name.lower() in rack_scope_values
        ]

    if section.target_type == "Device":
        devices = (
            session.execute(
                select(Device)
                .where(Device.project_id == project.id)
                .order_by(Device.id)
            )
            .scalars()
            .all()
        )
        filtered_devices: list[Device] = []
        for device in devices:
            if role_values and device.role.lower() not in role_values:
                continue
            if rack_scope_values:
                rack_id_value = (
                    str(device.rack_id).lower() if device.rack_id is not None else ""
                )
                rack_name_value = (
                    rack_id_to_name.get(device.rack_id, "").lower()
                    if device.rack_id is not None
                    else ""
                )
                if (
                    rack_id_value not in rack_scope_values
                    and rack_name_value not in rack_scope_values
                ):
                    continue
            filtered_devices.append(device)
        return filtered_devices

    return [project]


def _section_filters_allow(
    session: Session,
    section: SectionSnapshot,
    rule: SectionApplicationRule | None,
    project: Project,
) -> bool:
    targets = _target_objects_for_section(session, project, section, rule)
    if not targets:
        return False
    if rule is None or not rule.filters_json:
        return True
    return True


def _render_sections_for_project(
    session: Session, project_id: int
) -> list[dict[str, str]]:
    project = _load_project(session, project_id)
    snapshots = (
        session.execute(
            select(SectionSnapshot)
            .join(
                TemplateSetSnapshot,
                TemplateSetSnapshot.id == SectionSnapshot.template_set_snapshot_id,
            )
            .where(TemplateSetSnapshot.project_id == project_id)
            .order_by(
                SectionSnapshot.category.asc(),
                SectionSnapshot.section_order.asc(),
                SectionSnapshot.id.asc(),
            )
        )
        .scalars()
        .all()
    )
    rules = _section_rules_map(session, project_id)

    rendered: list[dict[str, str]] = []
    for section in snapshots:
        rule = rules.get(section.id)
        targets = _target_objects_for_section(session, project, section, rule)
        if not targets:
            continue

        for target in targets:
            context = {
                "project": project,
                "target": target,
                "now": datetime.now(UTC).isoformat(),
            }
            text = Template(section.text).render(**context)
            rendered.append(
                {
                    "category": section.category,
                    "target_type": section.target_type,
                    "output_targets": section.output_targets,
                    "text": text,
                }
            )
    return rendered


def _target_types_for_mode(mode: str) -> set[str]:
    if mode == "all":
        return {"word", "excel", "images"}
    if mode == "word":
        return {"word"}
    if mode == "excel":
        return {"excel"}
    if mode == "images":
        return {"images"}
    raise ValueError("Invalid mode")


def _operation_steps(session: Session, project_id: int) -> list[dict[str, str]]:
    location_map = _device_location_map(session, project_id)
    cablings = cablings_for_project(session, project_id)
    power_cablings = (
        session.execute(
            select(PowerCabling)
            .where(PowerCabling.project_id == project_id)
            .order_by(
                PowerCabling.label.asc(),
                PowerCabling.a_device.asc(),
                PowerCabling.a_port.asc(),
            )
        )
        .scalars()
        .all()
    )

    steps: list[dict[str, str]] = []
    steps.append(
        {
            "phase": "pre-check",
            "action": "Verify maintenance window, impact scope, and cable labels before touching any connection.",
            "expected_result": "Work permit and target endpoints are confirmed.",
            "rollback_hint": "Stop operation and escalate if scope is unclear.",
        }
    )

    for cabling in cablings:
        label = cabling.label or "(no-label)"
        steps.append(
            {
                "phase": "execution",
                "action": (
                    f"Connect network cable {label}: "
                    f"{cabling.a_device}:{cabling.a_port} -> {cabling.b_device}:{cabling.b_port} "
                    f"({cabling.cable_type or 'unspecified'}). "
                    f"Location: {location_map.get(cabling.a_device, 'unknown')} -> {location_map.get(cabling.b_device, 'unknown')}."
                ),
                "expected_result": "Endpoints are connected and cable label matches work order.",
                "rollback_hint": (
                    f"Restore previous endpoint mapping for {label} if link/error check fails."
                ),
            }
        )

    for power in power_cablings:
        label = power.label or "(no-label)"
        steps.append(
            {
                "phase": "execution",
                "action": (
                    f"Verify power path {label}: "
                    f"{power.a_device}:{power.a_port} -> {power.b_device}:{power.b_port} "
                    f"(bank={power.bank}, outlet={power.outlet}). "
                    f"Location: {location_map.get(power.a_device, 'unknown')} -> {location_map.get(power.b_device, 'unknown')}."
                ),
                "expected_result": "Power feed mapping is verified against plan.",
                "rollback_hint": "Revert to previous known-safe power outlet if mismatch is detected.",
            }
        )

    steps.append(
        {
            "phase": "post-check",
            "action": "Run link and service checks, then record completion evidence.",
            "expected_result": "All affected links/services are healthy and checklist is signed.",
            "rollback_hint": "Execute rollback procedure if service validation fails.",
        }
    )
    return steps


def _device_location_map(session: Session, project_id: int) -> dict[str, str]:
    devices = (
        session.execute(
            select(Device)
            .where(Device.project_id == project_id)
            .order_by(Device.name.asc(), Device.id.asc())
        )
        .scalars()
        .all()
    )
    racks = (
        session.execute(
            select(Rack).where(Rack.project_id == project_id).order_by(Rack.id.asc())
        )
        .scalars()
        .all()
    )
    rows = (
        session.execute(
            select(Row)
            .join(Room, Room.id == Row.room_id)
            .join(Site, Site.id == Room.site_id)
            .where(Site.project_id == project_id)
            .order_by(Row.id.asc())
        )
        .scalars()
        .all()
    )
    rooms = (
        session.execute(
            select(Room)
            .join(Site, Site.id == Room.site_id)
            .where(Site.project_id == project_id)
            .order_by(Room.id.asc())
        )
        .scalars()
        .all()
    )
    sites = (
        session.execute(
            select(Site)
            .where(Site.project_id == project_id)
            .order_by(Site.id.asc())
        )
        .scalars()
        .all()
    )

    racks_by_id = {rack.id: rack for rack in racks}
    rows_by_id = {row.id: row for row in rows}
    rooms_by_id = {room.id: room for room in rooms}
    sites_by_id = {site.id: site for site in sites}

    result: dict[str, str] = {}
    for device in devices:
        if device.rack_id is None:
            result[device.name] = "unassigned"
            continue

        rack = racks_by_id.get(device.rack_id)
        if rack is None:
            result[device.name] = "unassigned"
            continue

        parts = [f"rack={rack.name}"]
        if rack.row_id is not None:
            row_obj = rows_by_id.get(rack.row_id)
            if row_obj is not None:
                parts.append(f"row={row_obj.name}")
                room = rooms_by_id.get(row_obj.room_id)
                if room is not None:
                    parts.append(f"room={room.name}")
                    site = sites_by_id.get(room.site_id)
                    if site is not None:
                        parts.append(f"site={site.name}")
        result[device.name] = ", ".join(parts)
    return result


def _write_word(
    path: Path,
    rendered_sections: list[dict[str, str]],
    operation_steps: list[dict[str, str | int]],
) -> None:
    document = Document()
    last_category: str | None = None
    for section in rendered_sections:
        targets = json.loads(section["output_targets"])
        if "word" not in targets:
            continue
        category = section["category"]
        if category != last_category:
            document.add_heading(category, level=1)
            last_category = category
        for line in section["text"].splitlines() or [section["text"]]:
            if line.strip():
                document.add_paragraph(line)

    document.add_heading("Field Execution Pack", level=1)
    document.add_heading("Preconditions and Safety", level=2)
    document.add_paragraph(
        "Confirm approved maintenance window, affected scope, and safety constraints before execution."
    )
    document.add_paragraph(
        "Do not proceed if endpoint mapping, work permit, or rollback path is unclear."
    )

    document.add_heading("Step-by-step Procedure", level=2)
    for step in operation_steps:
        step_no = int(step["step_no"])
        depends_on = step["depends_on_step_no"]
        dependency_suffix = ""
        if isinstance(depends_on, int) and depends_on > 0:
            dependency_suffix = f" (depends on step {depends_on})"
        document.add_paragraph(
            f"{step_no}. [{step['phase']}] {step['action']}{dependency_suffix}"
        )
        document.add_paragraph(f"   Expected: {step['expected_result']}")

    document.add_heading("Post-work Verification", level=2)
    document.add_paragraph(
        "Validate link status, service status, and record execution evidence for each completed step."
    )

    document.add_heading("Rollback Procedure", level=2)
    document.add_paragraph(
        "If verification fails, revert affected changes in reverse step order and escalate with evidence."
    )
    document.save(path)


def _operation_steps_from_rendered_sections(
    rendered_sections: list[dict[str, str]],
) -> list[dict[str, str]]:
    steps: list[dict[str, str]] = []
    for section in rendered_sections:
        targets = json.loads(section["output_targets"])
        if "excel" not in targets and "word" not in targets:
            continue
        text = (section["text"] or "").strip()
        if not text:
            continue
        steps.append(
            {
                "phase": "execution",
                "action": text,
                "expected_result": "Section content is applied and verified.",
                "rollback_hint": "Revert section-level change and re-validate.",
            }
        )
    if not steps:
        steps.append(
            {
                "phase": "execution",
                "action": "No template-driven section steps were generated.",
                "expected_result": "Operator confirms no additional template steps are required.",
                "rollback_hint": "N/A",
            }
        )
    return steps


def _combined_operation_steps(
    session: Session,
    project_id: int,
    rendered_sections: list[dict[str, str]],
) -> list[dict[str, str | int]]:
    field_steps = _operation_steps(session, project_id)
    template_steps = _operation_steps_from_rendered_sections(rendered_sections)
    combined_steps: list[dict[str, str | int]] = []

    for raw_step in [*field_steps, *template_steps]:
        step_no = len(combined_steps) + 1
        dependency = step_no - 1 if step_no > 1 else ""
        combined_steps.append(
            {
                "step_no": step_no,
                "depends_on_step_no": dependency,
                "phase": raw_step["phase"],
                "action": raw_step["action"],
                "expected_result": raw_step["expected_result"],
                "rollback_hint": raw_step["rollback_hint"],
            }
        )
    return combined_steps


def _write_excel(
    path: Path,
    session: Session,
    project_id: int,
    rendered_sections: list[dict[str, str]],
    operation_steps: list[dict[str, str | int]],
) -> None:
    wb = Workbook()
    ws_wiring = wb.active
    ws_wiring.title = "wiring"
    ws_wiring.append(
        ["a_device", "a_port", "b_device", "b_port", "cable_type", "label"]
    )
    for cabling in cablings_for_project(session, project_id):
        ws_wiring.append(
            [
                cabling.a_device,
                cabling.a_port,
                cabling.b_device,
                cabling.b_port,
                cabling.cable_type or "",
                cabling.label or "",
            ]
        )

    ws_labels = wb.create_sheet("labels")
    ws_labels.append(["device", "role", "rack_id", "ru_start", "ru_size"])
    devices = (
        session.execute(
            select(Device)
            .where(Device.project_id == project_id)
            .order_by(Device.name, Device.id)
        )
        .scalars()
        .all()
    )
    for device in devices:
        ws_labels.append(
            [
                device.name,
                device.role,
                device.rack_id or "",
                device.ru_start or "",
                device.ru_size or "",
            ]
        )

    ws_checklists = wb.create_sheet("checklists")
    ws_checklists.append(["category", "text"])
    for section in rendered_sections:
        targets = json.loads(section["output_targets"])
        if "excel" in targets:
            ws_checklists.append([section["category"], section["text"]])

    ws_work_steps = wb.create_sheet("work_steps")
    ws_work_steps.append(
        [
            "step_no",
            "depends_on_step_no",
            "phase",
            "action",
            "expected_result",
            "rollback_hint",
        ]
    )
    for step in operation_steps:
        ws_work_steps.append(
            [
                step["step_no"],
                step["depends_on_step_no"],
                step["phase"],
                step["action"],
                step["expected_result"],
                step["rollback_hint"],
            ]
        )

    ws_verification = wb.create_sheet("verification_checklist")
    ws_verification.append(
        ["item", "method", "result", "evidence", "notes", "step_no"]
    )
    ws_verification.append(
        [
            "Maintenance window and scope confirmed",
            "Review work permit and endpoint list",
            "PENDING",
            "",
            "",
            "",
        ]
    )
    ws_verification.append(
        [
            "Cabling labels and endpoints verified",
            "Cross-check label against wiring sheet",
            "PENDING",
            "",
            "",
            "",
        ]
    )
    ws_verification.append(
        [
            "Post-work service checks completed",
            "Ping/health checks and operator confirmation",
            "PENDING",
            "",
            "",
            "",
        ]
    )
    for step in operation_steps:
        ws_verification.append(
            [
                f"Step {step['step_no']} completed",
                f"Verify: {step['expected_result']}",
                "PENDING",
                "",
                "",
                step["step_no"],
            ]
        )

    ws_issue_log = wb.create_sheet("issue_log")
    ws_issue_log.append(
        ["timestamp", "step_no", "severity", "issue", "action_taken", "owner"]
    )

    wb.save(path)


def _write_image_placeholder(path: Path, project_id: int, version_number: int) -> None:
    svg = (
        '<svg xmlns="http://www.w3.org/2000/svg" width="800" height="220">'
        '<rect width="100%" height="100%" fill="#ffffff" stroke="#000000"/>'
        f'<text x="24" y="60" font-size="24">Rackwright Project {project_id}</text>'
        f'<text x="24" y="100" font-size="18">Version {version_number} layout placeholder</text>'
        "</svg>"
    )
    path.write_text(svg, encoding="utf-8")


def _artifact_base_path(
    project_id: int, version_number: int, artifact_type: str
) -> Path:
    return (
        _data_dir()
        / "projects"
        / str(project_id)
        / "versions"
        / str(version_number)
        / artifact_type
    )


def _record_artifact_file(
    session: Session, artifact_version_id: int, artifact_type: str, file_path: Path
) -> None:
    base = _data_dir()
    relative = str(file_path.relative_to(base))
    session.add(
        ArtifactFile(
            artifact_version_id=artifact_version_id,
            artifact_type=artifact_type,
            relative_path=relative,
        )
    )


def generate(
    session: Session,
    project_id: int,
    mode: str,
    base_version_id: int | None,
    remarks: str | None,
    *,
    force: bool = False,
) -> GenerationResult:
    if mode not in VALID_MODES:
        raise ValueError("Invalid generation mode")

    _load_project(session, project_id)
    fingerprint = _compute_fingerprint(session, project_id, mode, base_version_id)
    existing = (
        session.execute(
            select(ArtifactVersion).where(
                ArtifactVersion.project_id == project_id,
                ArtifactVersion.fingerprint == fingerprint,
            )
        )
        .scalars()
        .one_or_none()
    )
    if existing is not None and not force:
        return GenerationResult(artifact_version=existing, skipped=True)

    version_number = _next_version_number(session, project_id)
    if existing is not None and force:
        fingerprint = hashlib.sha256(
            f"{fingerprint}|forced|{version_number}".encode("utf-8")
        ).hexdigest()
    targets = _target_types_for_mode(mode)

    artifact_version = ArtifactVersion(
        project_id=project_id,
        version_number=version_number,
        mode=mode,
        fingerprint=fingerprint,
        remarks=remarks,
        success_word=False,
        success_excel=False,
        success_image=False,
        errors_json=None,
    )
    session.add(artifact_version)
    session.flush()

    errors: list[dict[str, Any]] = []
    rendered_sections = _render_sections_for_project(session, project_id)
    operation_steps = _combined_operation_steps(session, project_id, rendered_sections)

    with change_log_context(source="generate", mode=mode, version=version_number):
        if "word" in targets:
            try:
                output_dir = _artifact_base_path(project_id, version_number, "word")
                output_dir.mkdir(parents=True, exist_ok=True)
                file_path = output_dir / "document.docx"
                _write_word(file_path, rendered_sections, operation_steps)
                _record_artifact_file(session, artifact_version.id, "word", file_path)
                artifact_version.success_word = True
            except Exception as exc:
                errors.append(
                    {
                        "type": "word",
                        "message": str(exc),
                        "jump_target": {"kind": "section", "value": "word-render"},
                    }
                )

        if "excel" in targets:
            try:
                output_dir = _artifact_base_path(project_id, version_number, "excel")
                output_dir.mkdir(parents=True, exist_ok=True)
                file_path = output_dir / "document.xlsx"
                _write_excel(
                    file_path,
                    session,
                    project_id,
                    rendered_sections,
                    operation_steps,
                )
                _record_artifact_file(session, artifact_version.id, "excel", file_path)
                artifact_version.success_excel = True
            except Exception as exc:
                errors.append(
                    {
                        "type": "excel",
                        "message": str(exc),
                        "jump_target": {"kind": "section", "value": "excel-render"},
                    }
                )

        if "images" in targets:
            try:
                output_dir = _artifact_base_path(project_id, version_number, "images")
                output_dir.mkdir(parents=True, exist_ok=True)
                file_path = output_dir / "rack_layout.svg"
                _write_image_placeholder(file_path, project_id, version_number)
                _record_artifact_file(session, artifact_version.id, "images", file_path)
                artifact_version.success_image = True
            except Exception as exc:
                errors.append(
                    {
                        "type": "images",
                        "message": str(exc),
                        "jump_target": {"kind": "section", "value": "image-render"},
                    }
                )

    artifact_version.errors_json = (
        json.dumps(errors, ensure_ascii=False) if errors else None
    )
    session.flush()
    return GenerationResult(artifact_version=artifact_version, skipped=False)

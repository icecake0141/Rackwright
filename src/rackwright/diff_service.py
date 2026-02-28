"""
Copyright 2026 Rackwright Contributors
SPDX-License-Identifier: Apache-2.0

This file was created or modified with the assistance of an AI (Large Language Model).
Review required for correctness, security, and licensing.
"""

from __future__ import annotations

import json
import os
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from openpyxl import load_workbook
from openpyxl.comments import Comment
from openpyxl.styles import PatternFill
from sqlalchemy import select
from sqlalchemy.orm import Session

from rackwright.models import ArtifactFile, ArtifactVersion, DiffItem, DiffReport


@dataclass
class DiffResult:
    diff_report: DiffReport


def _data_dir() -> Path:
    return Path(os.environ.get("RACKWRIGHT_DATA_DIR", "./data"))


def _artifact_path(
    session: Session, version_id: int, artifact_type: str
) -> Path | None:
    file_row = (
        session.execute(
            select(ArtifactFile).where(
                ArtifactFile.artifact_version_id == version_id,
                ArtifactFile.artifact_type == artifact_type,
            )
        )
        .scalars()
        .first()
    )
    if file_row is None:
        return None
    return _data_dir() / file_row.relative_path


def _artifact_output_base(
    project_id: int, version_number: int, artifact_type: str
) -> Path:
    return (
        _data_dir()
        / "projects"
        / str(project_id)
        / "versions"
        / str(version_number)
        / artifact_type
    )


def _record_artifact_file(
    session: Session, artifact_version_id: int, artifact_type: str, file_path: Path
) -> None:
    session.add(
        ArtifactFile(
            artifact_version_id=artifact_version_id,
            artifact_type=artifact_type,
            relative_path=str(file_path.relative_to(_data_dir())),
        )
    )


def _cell_value(value) -> str:
    if value is None:
        return ""
    return str(value)


def _diff_excel(
    from_path: Path, to_path: Path, output_path: Path
) -> list[dict[str, str]]:
    wb_from = load_workbook(from_path)
    wb_to = load_workbook(to_path)
    marked_wb = load_workbook(to_path)
    changed_items: list[dict[str, str]] = []

    for sheet_name in wb_to.sheetnames:
        ws_to = wb_to[sheet_name]
        ws_marked = marked_wb[sheet_name]
        ws_from = wb_from[sheet_name] if sheet_name in wb_from.sheetnames else None

        max_row = ws_to.max_row
        max_col = ws_to.max_column
        for row in range(1, max_row + 1):
            for col in range(1, max_col + 1):
                after_value = _cell_value(ws_to.cell(row=row, column=col).value)
                before_value = ""
                if (
                    ws_from is not None
                    and row <= ws_from.max_row
                    and col <= ws_from.max_column
                ):
                    before_value = _cell_value(ws_from.cell(row=row, column=col).value)
                if before_value != after_value:
                    cell = ws_marked.cell(row=row, column=col)
                    cell.fill = PatternFill(
                        start_color="FFF59D", end_color="FFF59D", fill_type="solid"
                    )
                    cell.comment = Comment(
                        f"Changed: {before_value} -> {after_value}", "rackwright"
                    )
                    changed_items.append(
                        {
                            "artifact_type": "excel",
                            "location": f"{sheet_name}!R{row}C{col}",
                            "change_type": "modified",
                            "before_value": before_value,
                            "after_value": after_value,
                        }
                    )

    diff_ws = marked_wb.create_sheet("diff_list")
    diff_ws.append(["sheet", "row", "col", "before", "after"])
    for item in changed_items:
        location = item["location"]
        sheet, rc = location.split("!")
        r_part, c_part = rc.split("C")
        row = int(r_part.replace("R", ""))
        col = int(c_part)
        diff_ws.append([sheet, row, col, item["before_value"], item["after_value"]])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    marked_wb.save(output_path)
    return changed_items


def _diff_word(
    from_path: Path, to_path: Path, output_path: Path
) -> list[dict[str, str]]:
    from_doc = Document(from_path)
    to_doc = Document(to_path)

    from_texts = [p.text for p in from_doc.paragraphs if p.text.strip()]
    to_texts = [p.text for p in to_doc.paragraphs if p.text.strip()]

    changed_items: list[dict[str, str]] = []
    max_len = max(len(from_texts), len(to_texts))
    for idx in range(max_len):
        before = from_texts[idx] if idx < len(from_texts) else ""
        after = to_texts[idx] if idx < len(to_texts) else ""
        if before != after:
            changed_items.append(
                {
                    "artifact_type": "word",
                    "location": f"paragraph:{idx+1}",
                    "change_type": "modified",
                    "before_value": before,
                    "after_value": after,
                }
            )

    for item in changed_items:
        location = item["location"]
        paragraph_index = int(location.split(":")[1]) - 1
        if 0 <= paragraph_index < len(to_doc.paragraphs):
            paragraph = to_doc.paragraphs[paragraph_index]
            if paragraph.runs:
                target_run = paragraph.runs[0]
            else:
                target_run = paragraph.add_run(
                    paragraph.text if paragraph.text else " "
                )
            target_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            to_doc.add_comment(
                runs=[target_run],
                text=f"Changed from: {item['before_value']}",
                author="rackwright",
                initials="RW",
            )

    to_doc.add_page_break()
    to_doc.add_heading("Word Diff Summary", level=1)
    to_doc.add_paragraph(f"Changed paragraphs: {len(changed_items)}")
    for item in changed_items:
        p = to_doc.add_paragraph()
        run = p.add_run(
            f"{item['location']}\nBEFORE: {item['before_value']}\nAFTER: {item['after_value']}"
        )
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    output_path.parent.mkdir(parents=True, exist_ok=True)
    to_doc.save(output_path)
    return changed_items


def diff(session: Session, from_version_id: int, to_version_id: int) -> DiffResult:
    from_version = session.get(ArtifactVersion, from_version_id)
    to_version = session.get(ArtifactVersion, to_version_id)
    if from_version is None or to_version is None:
        raise ValueError("Artifact version not found")
    if from_version.project_id != to_version.project_id:
        raise ValueError("Versions must belong to same project")

    existing = (
        session.execute(
            select(DiffReport).where(
                DiffReport.project_id == to_version.project_id,
                DiffReport.from_version_id == from_version_id,
                DiffReport.to_version_id == to_version_id,
            )
        )
        .scalars()
        .one_or_none()
    )
    if existing is not None:
        return DiffResult(diff_report=existing)

    diff_report = DiffReport(
        project_id=to_version.project_id,
        from_version_id=from_version_id,
        to_version_id=to_version_id,
        summary=None,
    )
    session.add(diff_report)
    session.flush()

    all_items: list[dict[str, str]] = []

    from_excel = _artifact_path(session, from_version_id, "excel")
    to_excel = _artifact_path(session, to_version_id, "excel")
    if from_excel is not None and to_excel is not None:
        out_excel = (
            _artifact_output_base(
                to_version.project_id, to_version.version_number, "diff"
            )
            / "excel_diff_list.xlsx"
        )
        excel_items = _diff_excel(from_excel, to_excel, out_excel)
        all_items.extend(excel_items)
        _record_artifact_file(session, to_version.id, "diff_excel", out_excel)

    from_word = _artifact_path(session, from_version_id, "word")
    to_word = _artifact_path(session, to_version_id, "word")
    if from_word is not None and to_word is not None:
        out_word = (
            _artifact_output_base(
                to_version.project_id, to_version.version_number, "diff"
            )
            / "word_diff_summary.docx"
        )
        word_items = _diff_word(from_word, to_word, out_word)
        all_items.extend(word_items)
        _record_artifact_file(session, to_version.id, "diff_word", out_word)

    for item in all_items:
        session.add(
            DiffItem(
                diff_report_id=diff_report.id,
                artifact_type=item["artifact_type"],
                location=item["location"],
                change_type=item["change_type"],
                before_value=item["before_value"],
                after_value=item["after_value"],
            )
        )

    summary = {
        "total_changes": len(all_items),
        "excel_changes": len([x for x in all_items if x["artifact_type"] == "excel"]),
        "word_changes": len([x for x in all_items if x["artifact_type"] == "word"]),
    }
    diff_report.summary = json.dumps(summary, ensure_ascii=False)
    session.flush()
    return DiffResult(diff_report=diff_report)
